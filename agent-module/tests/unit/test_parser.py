from io import BytesIO

import pytest
from docx import Document

from app.pipeline.parsers import (
    compute_content_hash,
    parse_docx_bytes_to_text,
    parse_email_html_to_text,
    parse_pdf_bytes_to_text,
)


def test_parse_email_html_to_text() -> None:
    html = "<html><body><h1>Hello</h1><p>Submit report by Friday</p></body></html>"
    text = parse_email_html_to_text(html)
    assert "Hello" in text
    assert "Submit report by Friday" in text


def test_parse_docx_bytes_to_text() -> None:
    doc = Document()
    doc.add_paragraph("Line 1")
    doc.add_paragraph("Line 2")
    buf = BytesIO()
    doc.save(buf)
    text = parse_docx_bytes_to_text(buf.getvalue())
    assert text == "Line 1\nLine 2"


def test_parse_pdf_bytes_to_text_with_mock(monkeypatch: pytest.MonkeyPatch) -> None:
    class _FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def get_text(self, *_args, **_kwargs) -> str:
            return self._text

    class _FakeDoc:
        def __iter__(self):
            return iter([_FakePage("A"), _FakePage("B")])

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    def _fake_open(*_args, **_kwargs):
        return _FakeDoc()

    monkeypatch.setattr("app.pipeline.parsers.fitz.open", _fake_open)
    text = parse_pdf_bytes_to_text(b"%PDF-dummy")
    assert text == "A\nB"


def test_compute_content_hash_is_deterministic() -> None:
    assert compute_content_hash("abc") == compute_content_hash("abc")
    assert compute_content_hash(b"abc") == compute_content_hash("abc")
